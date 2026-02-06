import os
import fitz  # PyMuPDF
from docx import Document
from langchain.docstore.document import Document as LC_Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from config import DATA_DIR, CHUNK_SIZE, CHUNK_OVERLAP
from PIL import Image
import fitz  # PyMuPDF
import io
def load_txt_file(path):
    with open(path, "r", encoding="utf-8") as f:
        text = f.read()
    return text

def load_pdf_file(path):
    text = ""
    doc = fitz.open(path)
    for page in doc:
        text += page.get_text()
    return text

def load_docx_file(path):
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_all_documents():
    documents = []

    for filename in os.listdir(DATA_DIR):
        path = os.path.join(DATA_DIR, filename)
        if filename.endswith(".txt"):
            text = load_txt_file(path)
        elif filename.endswith(".pdf"):
            text = load_pdf_file(path)
        elif filename.endswith(".docx"):
            text = load_docx_file(path)
        else:
            continue

        if not text.strip():
            continue

        documents.append(LC_Document(page_content=text, metadata={"source": filename}))

    return documents

def split_documents(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
    )
    return splitter.split_documents(documents)


def load_and_split_documents():
    raw_docs = load_all_documents()
    return split_documents(raw_docs)

def extract_images_from_pdf(path):
    doc = fitz.open(path)
    images = []
    for i, page in enumerate(doc):
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            images.append((f"{path}_page{i}_img{img_index}", img_pil))
    return images
